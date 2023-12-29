import streamlit as st
from webui_pages.utils import *
from streamlit_chatbox import *
# from server.chat.process_uploaded_file import *
from datetime import datetime
import os
import tempfile
import docx
from docx.shared import Pt

from configs import (LLM_MODEL, TEMPERATURE, HISTORY_LEN, PROMPT_TEMPLATES,
                     DEFAULT_KNOWLEDGE_BASE, DEFAULT_SEARCH_ENGINE,LANGCHAIN_LLM_MODEL, logger, log_verbose)
from typing import List, Dict

DEFAULT_INPUT = "default"

chat_box = ChatBox(
    assistant_avatar=os.path.join(
        "img",
        "chatchat_icon_blue_square_v2.png"
    )
)
def get_messages_history(history_len: int, content_in_expander: bool = False) -> List[Dict]:
    '''
    返回消息历史。
    content_in_expander控制是否返回expander元素中的内容，一般导出的时候可以选上，传入LLM的history不需要
    '''

    def filter(msg):
        content = [x for x in msg["elements"] if x._output_method in ["markdown", "text"]]
        if not content_in_expander:
            content = [x for x in content if not x._in_expander]
        content = [x.content for x in content]

        return {
            "role": msg["role"],
            "content": "\n\n".join(content),
        }

    return chat_box.filter_history(history_len=history_len, filter=filter)


def get_default_llm_model(api: ApiRequest) -> (str, bool):
    '''
    从服务器上获取当前运行的LLM模型，如果本机配置的LLM_MODEL属于本地模型且在其中，则优先返回
    返回类型为（model_name, is_local_model）
    '''
    running_models = api.list_running_models()
    if not running_models:
        return "", False

    if LLM_MODEL in running_models:
        return LLM_MODEL, True

    local_models = [k for k, v in running_models.items() if not v.get("online_api")]
    if local_models:
        return local_models[0], True
    return list(running_models)[0], False


def dialogue_page(api: ApiRequest):
    text = ""
    if not chat_box.chat_inited:
        default_model = get_default_llm_model(api)[0]
        st.toast(
            f"欢迎使用 [`Case-KnownledgeBase-glm`](https://github.com/chatchat-space/Langchain) ! \n\n"
            f"当前运行的模型`{default_model}`, 您可以开始提问了."
        )
        chat_box.init_session()
    with st.sidebar:
        # TODO: 对话模型与会话绑定
        def on_mode_change():
            mode = st.session_state.dialogue_mode
            text = f"已切换到 {mode} 模式。"
            if mode == "知识库问答":
                cur_kb = st.session_state.get("selected_kb")
                if cur_kb:
                    text = f"{text} 当前知识库： `{cur_kb}`。"
            st.toast(text)

        dialogue_mode = st.selectbox("请选择对话模式：",
                                     ["知识库问答",
                                      "LLM 对话",
                                      "搜索引擎问答",
                                      "自定义Agent问答",
                                      "LLM 摘要",
                                      ],
                                     index=0,
                                     on_change=on_mode_change,
                                     key="dialogue_mode",
                                     )

        def on_llm_change():
            if llm_model:
                config = api.get_model_config(llm_model)
                if not config.get("online_api"):  # 只有本地model_worker可以切换模型
                    st.session_state["prev_llm_model"] = llm_model
                st.session_state["cur_llm_model"] = st.session_state.llm_model

        def llm_model_format_func(x):
            if x in running_models:
                return f"{x} (Running)"
            return x

        running_models = list(api.list_running_models())
        running_models += LANGCHAIN_LLM_MODEL.keys()
        available_models = []
        config_models = api.list_config_models()
        worker_models = list(config_models.get("worker", {}))  # 仅列出在FSCHAT_MODEL_WORKERS中配置的模型
        for m in worker_models:
            if m not in running_models and m != "default":
                available_models.append(m)
        for k, v in config_models.get("online", {}).items():  # 列出ONLINE_MODELS中直接访问的模型
            if not v.get("provider") and k not in running_models:
                available_models.append(k)
        for k, v in config_models.get("langchain", {}).items():  # 列出LANGCHAIN_LLM_MODEL支持的模型
            available_models.append(k)
        llm_models = running_models + available_models
        index = llm_models.index(st.session_state.get("cur_llm_model", get_default_llm_model(api)[0]))
        llm_model = st.selectbox("选择LLM模型：",
                                 llm_models,
                                 index,
                                 format_func=llm_model_format_func,
                                 on_change=on_llm_change,
                                 key="llm_model",
                                 )
        if (st.session_state.get("prev_llm_model") != llm_model
                and not llm_model in config_models.get("online", {})
                and not llm_model in config_models.get("langchain", {})
                and llm_model not in running_models):
            with st.spinner(f"正在加载模型： {llm_model}，请勿进行操作或刷新页面"):
                prev_model = st.session_state.get("prev_llm_model")
                r = api.change_llm_model(prev_model, llm_model)
                if msg := check_error_msg(r):
                    st.error(msg)
                elif msg := check_success_msg(r):
                    st.success(msg)
                    st.session_state["prev_llm_model"] = llm_model

        index_prompt = {
            "LLM 对话": "llm_chat",
            "自定义Agent问答": "agent_chat",
            "搜索引擎问答": "search_engine_chat",
            "知识库问答": "knowledge_base_chat",
            "LLM 摘要": "llm_abstract",
        }
        prompt_templates_kb_list = list(PROMPT_TEMPLATES[index_prompt[dialogue_mode]].keys())
        prompt_template_name = prompt_templates_kb_list[0]
        if "prompt_template_select" not in st.session_state:
            st.session_state.prompt_template_select = prompt_templates_kb_list[0]

        def prompt_change():
            text = f"已切换为 {prompt_template_name} 模板。"
            st.toast(text)

        prompt_template_select = st.selectbox(
            "请选择Prompt模板：",
            prompt_templates_kb_list,
            index=0,
            on_change=prompt_change,
            key="prompt_template_select",
        )
        prompt_template_name = st.session_state.prompt_template_select
        temperature = st.slider("Temperature：", 0.0, 1.0, TEMPERATURE, 0.05)
        history_len = st.number_input("历史对话轮数：", 0, 20, HISTORY_LEN)

        def on_kb_change():
            st.toast(f"已加载知识库： {st.session_state.selected_kb}")
        

        if dialogue_mode == "知识库问答":
            with st.expander("知识库配置", True):
                kb_list = api.list_knowledge_bases()
                index = 0
                if DEFAULT_KNOWLEDGE_BASE in kb_list:
                    index = kb_list.index(DEFAULT_KNOWLEDGE_BASE)
                selected_kb = st.selectbox(
                    "请选择知识库：",
                    kb_list,
                    index=index,
                    on_change=on_kb_change,
                    key="selected_kb",
                )
                kb_top_k = st.number_input("匹配知识条数：", 1, 20, VECTOR_SEARCH_TOP_K)

                ## Bge 模型会超过1
                score_threshold = st.slider("知识匹配分数阈值：", 0.0, 1.0, float(SCORE_THRESHOLD), 0.01)

        elif dialogue_mode == "搜索引擎问答":
            search_engine_list = api.list_search_engines()
            if DEFAULT_SEARCH_ENGINE in search_engine_list:
                index = search_engine_list.index(DEFAULT_SEARCH_ENGINE)
            else:
                index = search_engine_list.index("duckduckgo") if "duckduckgo" in search_engine_list else 0
            with st.expander("搜索引擎配置", True):
                search_engine = st.selectbox(
                    label="请选择搜索引擎",
                    options=search_engine_list,
                    index=index,
                )
                se_top_k = st.number_input("匹配搜索结果条数：", 1, 20, SEARCH_ENGINE_TOP_K)
        
        # 通过文件上传按钮获取文件内容，如果用户上传了文件，将文件内容作为 prompt

        uploaded_file = st.file_uploader("请选择输入文档", key="file_uploader")
        # if uploaded_file is not None:
        #     logger.info(uploaded_file)
        #     logger.info(uploaded_file.read().decode("utf-8"))
        confirm_run_button = st.button("确认运行")


    # Display chat messages from history on app rerun
    chat_box.output_messages()
  

    chat_input_placeholder = "请输入对话内容，换行请使用Shift+Enter "


    # If the "确认运行" button is not clicked or no file is uploaded, process chat input
    if prompt := st.chat_input(chat_input_placeholder, key="prompt"):
        history = get_messages_history(history_len)
        chat_box.user_say(prompt)

        if dialogue_mode == "LLM 对话":
            chat_box.ai_say("正在思考...")
            text = ""
            r = api.chat_chat(prompt,
                            history=history,
                            model=llm_model,
                            prompt_name=prompt_template_name,
                            temperature=temperature)
            for t in r:
                if error_msg := check_error_msg(t):  # check whether error occured
                    st.error(error_msg)
                    break
                text += t
                chat_box.update_msg(text)
            chat_box.update_msg(text, streaming=False)  # 更新最终的字符串，去除光标



        elif dialogue_mode == "自定义Agent问答":
            chat_box.ai_say([
                f"正在思考...",
                Markdown("...", in_expander=True, title="思考过程", state="complete"),

            ])
            text = ""
            ans = ""
            support_agent = ["Azure-OpenAI", "OpenAI", "Anthropic", "Qwen", "qwen-api", "baichuan-api"]  # 目前支持agent的模型
            if not any(agent in llm_model for agent in support_agent):
                ans += "正在思考... \n\n <span style='color:red'>该模型并没有进行Agent对齐，请更换支持Agent的模型获得更好的体验！</span>\n\n\n"
                chat_box.update_msg(ans, element_index=0, streaming=False)
            for d in api.agent_chat(prompt,
                                    history=history,
                                    model=llm_model,
                                    prompt_name=prompt_template_name,
                                    temperature=temperature,
                                    ):
                try:
                    d = json.loads(d)
                except:
                    pass
                if error_msg := check_error_msg(d):  # check whether error occured
                    st.error(error_msg)
                if chunk := d.get("answer"):
                    text += chunk
                    chat_box.update_msg(text, element_index=1)
                if chunk := d.get("final_answer"):
                    ans += chunk
                    chat_box.update_msg(ans, element_index=0)
                if chunk := d.get("tools"):
                    text += "\n\n".join(d.get("tools", []))
                    chat_box.update_msg(text, element_index=1)
            chat_box.update_msg(ans, element_index=0, streaming=False)
            chat_box.update_msg(text, element_index=1, streaming=False)
        elif dialogue_mode == "知识库问答":
            chat_box.ai_say([
                f"正在查询知识库 `{selected_kb}` ...",
                Markdown("...", in_expander=True, title="知识库匹配结果", state="complete"),
            ])
            text = ""
            for d in api.knowledge_base_chat(prompt,
                                            knowledge_base_name=selected_kb,
                                            top_k=kb_top_k,
                                            score_threshold=score_threshold,
                                            history=history,
                                            model=llm_model,
                                            prompt_name=prompt_template_name,
                                            temperature=temperature):
                if error_msg := check_error_msg(d):  # check whether error occured
                    st.error(error_msg)
                elif chunk := d.get("answer"):
                    text += chunk
                    chat_box.update_msg(text, element_index=0)
            chat_box.update_msg(text, element_index=0, streaming=False)
            chat_box.update_msg("\n\n".join(d.get("docs", [])), element_index=1, streaming=False)

        elif dialogue_mode == "搜索引擎问答":
            chat_box.ai_say([
                f"正在执行 `{search_engine}` 搜索...",
                Markdown("...", in_expander=True, title="网络搜索结果", state="complete"),
            ])
            text = ""
            for d in api.search_engine_chat(prompt,
                                            search_engine_name=search_engine,
                                            top_k=se_top_k,
                                            history=history,
                                            model=llm_model,
                                            prompt_name=prompt_template_name,
                                            temperature=temperature):
                if error_msg := check_error_msg(d):  # check whether error occured
                    st.error(error_msg)
                elif chunk := d.get("answer"):
                    text += chunk
                    chat_box.update_msg(text, element_index=0)
            chat_box.update_msg(text, element_index=0, streaming=False)
            chat_box.update_msg("\n\n".join(d.get("docs", [])), element_index=1, streaming=False)

        elif dialogue_mode == "LLM 摘要":
            chat_box.ai_say("正在进行摘要...")
            text = ""
            r = api.chat_abstract(prompt,
                            history=history,
                            model=llm_model,
                            prompt_name=prompt_template_name,
                            temperature=temperature)
            for t in r:
                if error_msg := check_error_msg(t):  # check whether error occured
                    st.error(error_msg)
                    break
                text += t
                chat_box.update_msg(text)
            chat_box.update_msg(text, streaming=False)  # 更新最终的字符串，去除光标
    
    #文档处理部分
    # 在调用 process_uploaded_file 之前获取部分知识库相关值
    # selected_kb = st.session_state.get("selected_kb")
    # kb_top_k = st.session_state.get("kb_top_k")
    # score_threshold = st.session_state.get("score_threshold")
    # search_engine = st.session_state.get("search_engine")
    # se_top_k = st.session_state.get("se_top_k")
    # selected_kb = "test1"
    # kb_top_k = 1
    # score_threshold = 0.55
    # search_engine = "duckduckgo"
    # se_top_k = 1


    # if confirm_run_button and uploaded_file is not None:
    #     text = api.process_uploaded_file(uploaded_file, dialogue_mode, llm_model, prompt_template_name, temperature, history_len, selected_kb, kb_top_k, score_threshold, search_engine, se_top_k)
    #     print(f"received output message:", text)
            
    if confirm_run_button and uploaded_file is not None:
        # read txt
        # prompt = uploaded_file.read().decode('utf-8')

        # read docx
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(uploaded_file.read())
            logger.info(f"文件路径为：{temp_file.name}")

        doc = docx.Document(temp_file.name)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        prompt = content

        # delete tempfile
        temp_file_path = temp_file.name
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

        prompt_length = len(prompt)
        history = get_messages_history(history_len)
        chat_box.user_say(f"用户输入文档长度为: {prompt_length}")
        # chat_box.ai_say("正确运行")
        if dialogue_mode == "LLM 对话":
            chat_box.ai_say("正在思考...")
            text = ""
            r = api.chat_chat(prompt,
                            history=history,
                            model=llm_model,
                            prompt_name=prompt_template_name,
                            temperature=temperature)
            for t in r:
                if error_msg := check_error_msg(t):  # check whether error occured
                    st.error(error_msg)
                    break
                text += t
                chat_box.update_msg(text)
            chat_box.update_msg(text, streaming=False)  # 更新最终的字符串，去除光标



        elif dialogue_mode == "自定义Agent问答":
            chat_box.ai_say([
                f"正在思考...",
                Markdown("...", in_expander=True, title="思考过程", state="complete"),

            ])
            text = ""
            ans = ""
            support_agent = ["Azure-OpenAI", "OpenAI", "Anthropic", "Qwen", "qwen-api", "baichuan-api"]  # 目前支持agent的模型
            if not any(agent in llm_model for agent in support_agent):
                ans += "正在思考... \n\n <span style='color:red'>该模型并没有进行Agent对齐，请更换支持Agent的模型获得更好的体验！</span>\n\n\n"
                chat_box.update_msg(ans, element_index=0, streaming=False)
            for d in api.agent_chat(prompt,
                                    history=history,
                                    model=llm_model,
                                    prompt_name=prompt_template_name,
                                    temperature=temperature,
                                    ):
                try:
                    d = json.loads(d)
                except:
                    pass
                if error_msg := check_error_msg(d):  # check whether error occured
                    st.error(error_msg)
                if chunk := d.get("answer"):
                    text += chunk
                    chat_box.update_msg(text, element_index=1)
                if chunk := d.get("final_answer"):
                    ans += chunk
                    chat_box.update_msg(ans, element_index=0)
                if chunk := d.get("tools"):
                    text += "\n\n".join(d.get("tools", []))
                    chat_box.update_msg(text, element_index=1)
            chat_box.update_msg(ans, element_index=0, streaming=False)
            chat_box.update_msg(text, element_index=1, streaming=False)

        elif dialogue_mode == "知识库问答":
            chat_box.ai_say([
                f"正在查询知识库 `{selected_kb}` ...",
                Markdown("...", in_expander=True, title="知识库匹配结果", state="complete"),
            ])
            abstract = ""
            text = ""
            r = api.chat_abstract(prompt,
                            history=history,
                            model=llm_model,
                            prompt_name="test2-b",
                            temperature=temperature)
            for t in r:
                if error_msg := check_error_msg(t):  # check whether error occured
                    st.error(error_msg)
                    break
                abstract += t
            #     chat_box.update_msg(abstract)
            # chat_box.update_msg(abstract, streaming=False)  # 更新最终的字符串，去除光标
            
            # Assuming your abstract points are separated by newline characters ("\n")
            abstract_points = abstract.split("\n")
            for point in abstract_points:
                point = point.strip()  # 去除前导和尾随的空格
                if not point:
                    continue  # 跳过空行
                for d in api.knowledge_base_chat(point,
                                                knowledge_base_name=selected_kb,
                                                top_k=kb_top_k,
                                                score_threshold=score_threshold,
                                                history=history,
                                                model=llm_model,
                                                prompt_name=prompt_template_name,
                                                temperature=temperature):
                    if error_msg := check_error_msg(d):  # check whether error occured
                        st.error(error_msg)
                    elif chunk := d.get("answer"):
                        text += chunk
                        chat_box.update_msg(text, element_index=0)
            chat_box.update_msg(text, element_index=0, streaming=False)
            chat_box.update_msg("\n\n".join(d.get("docs", [])), element_index=1, streaming=False)

        elif dialogue_mode == "搜索引擎问答":
            chat_box.ai_say([
                f"正在执行 `{search_engine}` 搜索...",
                Markdown("...", in_expander=True, title="网络搜索结果", state="complete"),
            ])
            text = ""
            for d in api.search_engine_chat(prompt,
                                            search_engine_name=search_engine,
                                            top_k=se_top_k,
                                            history=history,
                                            model=llm_model,
                                            prompt_name=prompt_template_name,
                                            temperature=temperature):
                if error_msg := check_error_msg(d):  # check whether error occured
                    st.error(error_msg)
                elif chunk := d.get("answer"):
                    text += chunk
                    chat_box.update_msg(text, element_index=0)
            chat_box.update_msg(text, element_index=0, streaming=False)
            chat_box.update_msg("\n\n".join(d.get("docs", [])), element_index=1, streaming=False)
    
        elif dialogue_mode == "LLM 摘要":
            chat_box.ai_say("正在进行摘要...")
            text = ""
            r = api.chat_abstract(prompt,
                            history=history,
                            model=llm_model,
                            prompt_name=prompt_template_name,
                            temperature=temperature)
            for t in r:
                if error_msg := check_error_msg(t):  # check whether error occured
                    st.error(error_msg)
                    break
                text += t
                chat_box.update_msg(text)
            chat_box.update_msg(text, streaming=False)  # 更新最终的字符串，去除光标

            
    now = datetime.now()

    # 获取最后一次聊天历史的文本表示
    def get_last_chat_history_text(chat_box):
        last_msg = chat_box.history[-1] if chat_box.history else None
        if last_msg:
            content = "\n".join([element.content for element in last_msg["elements"]])
            return f"{content}\n\n"
        else:
            return "No chat history available."
    
    chat_history_text = get_last_chat_history_text(chat_box)
    logger.info(chat_history_text)
    # save text to docx
    doc = docx.Document()
    paragraph = doc.add_paragraph()

    run = paragraph.add_run(chat_history_text)
    font = run.font
    # font.name = 'SimSun' 
    font.size = Pt(12)

    output_path = "output.docx"
    doc.save(output_path)
    
    with st.sidebar:
        
        cols = st.columns(2)
        export_btn = cols[0]
        if cols[1].button(
                "清空对话",
                use_container_width=True,
        ):
            chat_box.reset_history()
            st.experimental_rerun()

        export_btn.download_button(
            "导出记录",
            "".join(chat_box.export2md()),
            file_name=f"{now:%Y-%m-%d %H.%M}_对话记录.md",
            mime="text/markdown",
            use_container_width=True,
        )

    st.download_button(
            "下载txt结果",
            get_last_chat_history_text(chat_box),
            file_name=f"{now:%Y-%m-%d %H.%M}_结果文件.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with open(output_path, "rb") as file:
        st.download_button(
                "下载docx结果",
                file,
                file_name=f"{now:%Y-%m-%d %H.%M}_结果文件.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )