import streamlit as st
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from configs import logger
from webui_pages.utils import *
from server.utils import api_address
import docx
from docx.shared import Pt
from datetime import datetime
import tempfile
import platform


async def test_upload(file: UploadFile = File(...)):
    api = AsyncApiRequest(base_url=api_address())
    operating_system = platform.system()
    if operating_system == 'Windows':
        logger.info("电脑操作系统是 Windows")
    elif operating_system == 'Linux':
        logger.info("电脑操作系统是 Linux")
    else:
        logger.info("未知操作系统")

    # read text(.txt)
    # content = await file.read()
    # prompt = content.decode("utf-8")

    # read docx
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        temp_file.write(file.file.read())
        logger.info(f"文件路径为：{temp_file.name}")

    doc = docx.Document(temp_file.name)
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"
    prompt = content

    # delete tempfile
    temp_file_path = temp_file.name
    if os.path.exists(temp_file_path):
        os.remove(temp_file_path)

    #logger.info(f"文件路径为：{temp_file.name}")

    # chat function
    logger.info(f"获取到的输入为：{prompt}")
    text = ""
    r = api.chat_chat(prompt,
                    model="chatglm3-6b",
                    prompt_name="default",
                    temperature=0.5)
    async for t in r:
        if error_msg := check_error_msg(t):  # check whether error occured
            st.error(error_msg)
            break
        # logger.info(t)
        text += t

    # save text to docx
    doc = docx.Document()
    paragraph = doc.add_paragraph()

    run = paragraph.add_run(text)
    font = run.font
    # font.name = 'SimSun' 
    font.size = Pt(12)

    # 尝试使用临时文件，但return时临时文件已被删除
    # with tempfile.TemporaryDirectory() as temp_dir:
    #     output_path = os.path.join(temp_dir, "output.docx")
    #     logger.info(f"临时文件夹为：{temp_dir}")

    output_path = "output.docx"
    doc.save(output_path)
    now = datetime.now()
    # Return the docx file as a FileResponse
    return FileResponse(output_path, filename=f"{now:%Y-%m-%d %H.%M}_结果文件.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


