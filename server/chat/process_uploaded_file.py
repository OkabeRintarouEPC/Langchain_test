import streamlit as st
from fastapi import Body, File, UploadFile, Form
from configs import (LLM_MODEL, TEMPERATURE, HISTORY_LEN, SCORE_THRESHOLD, PROMPT_TEMPLATES,
                     DEFAULT_KNOWLEDGE_BASE, DEFAULT_SEARCH_ENGINE,LANGCHAIN_LLM_MODEL, logger, log_verbose)
from webui_pages.utils import *
from streamlit_chatbox import *
from datetime import datetime
import os
import random
from fastapi.responses import JSONResponse, FileResponse
import tempfile
import docx
from docx.shared import Pt
import platform


async def process_uploaded_file(
                        uploaded_file: UploadFile = File(..., description="上传docx文件，不支持多文件"),
                        llm_model: str = Body(LLM_MODEL, description="LLM 模型名称"),
                        prompt_template_name: str = Body("caseKL_2", description="使用的知识库问答prompt模板名称(在configs/prompt_config.py中配置)"), 
                        temperature: float = Body(TEMPERATURE, description="LLM 采样温度", ge=0.0, le=1.0), 
                        history_len: int = Body(HISTORY_LEN, description="历史对话轮数"), 
                        selected_kb: str = Body("test1", description="使用的知识库"), 
                        kb_top_k: int = Body(VECTOR_SEARCH_TOP_K, description="知识库匹配向量数量"), 
                        score_threshold: float = Body(SCORE_THRESHOLD, description="知识库匹配相关度阈值"), 
                        output_type: int = Body(0, description="输出格式，输出docx为0，输出json为1")
                    )-> JSONResponse:
    # 注意：在异步函数内部，使用异步ApiRequest
    api = AsyncApiRequest(base_url=api_address())
    operating_system = platform.system()

    # Read txt
    # content = await uploaded_file.read()
    # prompt = content.decode("utf-8")

    # Read docx
    if operating_system == 'Windows':
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(uploaded_file.file.read())
            #logger.info(f"文件路径为：{temp_file.name}")

        doc = docx.Document(temp_file.name)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        prompt = content

        # Delete tempfile
        temp_file_path = temp_file.name
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

    elif operating_system == 'Linux':
        with tempfile.NamedTemporaryFile(delete=True) as temp_file:
            temp_file.write(uploaded_file.file.read())
            #logger.info(f"文件路径为：{temp_file.name}")
            doc = docx.Document(temp_file.name)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            prompt = content

    else:
        error_message = "This platform is not currently supported."
        return JSONResponse(content={"error": error_message}, status_code=400)


    # 先进行文档摘要，再将摘要分点送入LLM进行知识库问答
    abstract = ""
    text = ""
    r = api.chat_abstract(prompt,
                    model=llm_model,
                    prompt_name="test2-b",
                    temperature=temperature)
    async for t in r:
        if error_msg := check_error_msg(t):  # check whether error occured
            st.error(error_msg)
            break
        abstract += t

    # Assuming the abstract points are separated by "\n"
    abstract_points = abstract.split("\n")
    for point in abstract_points:
        point = point.strip()  # Remove leading and trailing spaces
        if not point:
            continue  # Skip empty section
        async for d in api.knowledge_base_chat(point,
                                        knowledge_base_name=selected_kb,
                                        top_k=kb_top_k,
                                        score_threshold=score_threshold,
                                        model=llm_model,
                                        prompt_name=prompt_template_name,
                                        temperature=temperature):
            if error_msg := check_error_msg(d):  # check whether error occured
                st.error(error_msg)
            elif chunk := d.get("answer"):
                text += chunk
    
    # 根据平台和返回类型决定输出方式
    if operating_system == 'Windows':
        # Return as json
        if output_type == 1:
            return JSONResponse(content=text)
        # Return as docx
        else:
            doc = docx.Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)

            # Set docx display format
            font = run.font
            # font.name = 'SimSun' 
            font.size = Pt(12)

            random_suffix = random.randint(1000, 9999)
            now = datetime.now()
            output_path = f"{now:%Y-%m-%d %H.%M}_{random_suffix}_output.docx"
            doc.save(output_path)
            return FileResponse(output_path, filename=f"{now:%Y-%m-%d %H.%M}_结果文件.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
   
    elif operating_system == 'Linux':
        # Return as json
        if output_type == 1:
            return JSONResponse(content=text)
        # Return as docx
        else:
            doc = docx.Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)

            # Set docx display format
            font = run.font
            # font.name = 'SimSun' 
            font.size = Pt(12)

            # Try to use tempfile, maybe works on Linux
            with tempfile.TemporaryDirectory() as temp_dir:
                random_suffix = random.randint(1000, 9999)
                now = datetime.now()
                output_path = os.path.join(temp_dir, f"{now:%Y-%m-%d %H.%M}_{random_suffix}_output.docx")
                # logger.info(f"临时文件夹为：{temp_dir}")
                doc.save(output_path)
                return FileResponse(output_path, filename=f"{now:%Y-%m-%d %H.%M}_结果文件.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
  
    else:
        error_message = "This platform is not currently supported."
        return JSONResponse(content={"error": error_message}, status_code=400)


